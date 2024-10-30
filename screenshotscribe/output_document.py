from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import Pt

from screenshotscribe.utils import open_file


def output_document(generated_data):
    document = Document()
    style = document.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(12)
    
    status = document.add_paragraph()
    errors = 0

    for item in generated_data:
        bullet = document.add_paragraph(style='List Bullet')
        run = bullet.add_run(item['text'])
        if item['success'] != True:
            errors += 1
            run2 = bullet.add_run(' ' + item['filename'] + ' ' +
                                 item['date_created'])
            run.bold = True
            run2.bold = True
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            run2.font.highlight_color = WD_COLOR_INDEX.YELLOW
            continue

    if errors == 0:
        status_run = status.add_run(f'{len(generated_data)} images successfully transcribed')
    else:
        status_run = status.add_run(f'{errors}/{len(generated_data)} images not able to be transcribed. Review each corresponding finish reason for more information')
        status_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    status_run.bold = True

    filename = 'transcribed.docx'
    document.save(filename)
    print(f'Document generated at {filename}')

    open_file(filename)
