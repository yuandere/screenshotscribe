import datetime
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import RGBColor

document = Document()


def get_readable_time(ctime):
    readable = datetime.datetime.fromtimestamp(ctime)
    return readable.strftime('%m/%d/%Y %H:%M:%S')


def generate_document(generated_data):
    document.add_heading('formatting test', level=1)
    sh = document.add_paragraph()
    errors = 0

    for item in generated_data:
        bullet = document.add_paragraph(item['text'], style='List Bullet')
        if item['success'] != True:
            errors += 1
            run = bullet.add_run(item['filename'] + ' ' +
                                 get_readable_time(item['date_created']))
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            continue

    sh.text = f'{errors}/{len(generated_data)} images not able to be transcribed. Review each corresponding finish_reason for more information' if errors > 0 else f'{
        len(generated_data)} images successfully transcribed'

    document.save('transcribed.docx')
